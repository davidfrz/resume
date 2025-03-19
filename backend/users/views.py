# coding:utf-8
from django.shortcuts import render
from rest_framework import generics, permissions, status
from rest_framework.response import Response
from rest_framework.views import APIView
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.models import User
from .serializers import UserSerializer, RegisterSerializer
from rest_framework.parsers import FormParser, MultiPartParser, JSONParser
from docx import Document
from docx.shared import Inches
from pprint import pprint
from paddlenlp import Taskflow
import fitz  # 用于PDF转图片

import os
from openai import OpenAI
import json
from paddleocr import PaddleOCR
client = OpenAI(
    # 若没有配置环境变量，请用百炼API Key将下行替换为：api_key="sk-xxx",
    api_key="", 
    base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
)
ocr = PaddleOCR(use_angle_cls=True, lang="ch")
# 简历图片上传测试视图
class UploadImageView(APIView):
    parser_classes = [MultiPartParser, FormParser]
    permission_classes = [permissions.AllowAny]

    def post(self, request):
        # 获取上传的图片文件和实体类型
        image_file = request.FILES.get('file')
        entities = request.data.get('entities', '').split(',')  # 接收前端传来的实体类型

        if not image_file:
            return Response({
                'code': 400,
                'message': '未上传图片文件或文件格式不正确'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 验证文件类型
        allowed_types = ['image/jpeg', 'image/png', 'image/jpg']
        if image_file.content_type not in allowed_types:
            return Response({
                'code': 400,
                'message': '不支持的文件类型，请上传JPG或PNG格式的图片'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 生成唯一的文件名
        import os
        from datetime import datetime
        from django.conf import settings
        
        file_name = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}{os.path.splitext(image_file.name)[1]}"
        
        # 确保media目录存在
        os.makedirs(settings.MEDIA_ROOT, exist_ok=True)
        
        # 使用MEDIA_ROOT构建完整的文件路径
        file_path = os.path.join(settings.MEDIA_ROOT, file_name)
        
        # 保存文件
        with open(file_path, 'wb+') as destination:
            for chunk in image_file.chunks():
                destination.write(chunk)
        
        # 动态设置schema
        schema = [entity.strip() for entity in entities if entity.strip()]  # 去除空字符串
        if not schema:  # 如果用户没选择，默认使用这些
            schema = ['姓名', '出生日期', '电话']
        
        # ie = Taskflow('information_extraction', schema=schema)

        # result = ie({"doc": file_path})
        data = {}
        # for key in schema:
        #     if key in result[0]:
        #         data[key] = result[0][key][0]['text']
        
        
        mes = ocr.ocr(file_path, cls=True)
        mes_str = json.dumps(mes, ensure_ascii=False).replace("\n", "").replace(" ","")
        completion = client.chat.completions.create(
        model="qwen-plus", # 此处以qwen-plus为例，可按需更换模型名称。模型列表：https://help.aliyun.com/zh/model-studio/getting-started/models
        messages=[
            {'role': 'system', 'content': '你是一个简历提取专家，我会给你文字形式的简历，里面还有很多杂乱的数字，你只关心简历的部分，并把它给我提取成对应的json格式返回给我，例如：{姓名：xxx，年龄：xxx}.注意，请只给我简历的部分，不要别的任何东西'},
            {'role': 'user', 'content': mes_str}],
        )
        res = completion.model_dump_json()
        res = json.loads(res)
        cont = res["choices"][0]["message"]["content"]
        # print(cont)
        data = data = json.loads(cont)
        data['image_url'] = request.build_absolute_uri(f'/media/{file_name}')
        
        # 返回数据
        return Response({
            'code': 200,
            'data': data,
            'message': '解析成功'
        }, status=status.HTTP_200_OK)
# 用户注册视图
class RegisterView(generics.CreateAPIView):
    queryset = User.objects.all()
    permission_classes = [permissions.AllowAny]
    serializer_class = RegisterSerializer

# 用户登录视图
class LoginView(APIView):
    permission_classes = [permissions.AllowAny]
    parser_classes = [FormParser, MultiPartParser, JSONParser]
    
    def post(self, request):
        username = request.data.get('username')
        password = request.data.get('password')
        user = authenticate(username=username, password=password)
        print(password)
        if user is not None:
            login(request, user)
            serializer = UserSerializer(user)
            return Response(serializer.data)
        return Response({'detail': '用户名或密码错误'}, status=status.HTTP_400_BAD_REQUEST)

# 用户登出视图
class LogoutView(APIView):
    permission_classes = [permissions.AllowAny]
    def post(self, request):
        logout(request)
        return Response({'detail': '成功登出'})

# 获取当前用户信息
class UserDetailView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    
    def get(self, request):
        serializer = UserSerializer(request.user)
        return Response(serializer.data)


# PDF上传处理视图
class UploadPDFView(APIView):
    parser_classes = [MultiPartParser, FormParser]
    permission_classes = [permissions.AllowAny]

    def post(self, request):
        # 获取上传的 PDF 文件和实体类型
        pdf_file = request.FILES.get('file')
        entities = request.data.get('entities', '').split(',')  # 接收前端传来的实体类型

        if not pdf_file:
            return Response({
                'code': 400,
                'message': '未上传 PDF 文件或文件格式不正确'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 验证文件类型
        allowed_types = ['application/pdf']
        if pdf_file.content_type not in allowed_types:
            return Response({
                'code': 400,
                'message': '不支持的文件类型，请上传 PDF 格式的文件'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 生成唯一的文件名
        import os
        from datetime import datetime
        from django.conf import settings
        
        # file_name = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}{os.path.splitext(pdf_file.name)[1]}"
        file_name = 'test_pdf.pdf'
        
        # 确保 media 目录存在
        os.makedirs(settings.MEDIA_ROOT, exist_ok=True)
        
        # 使用 MEDIA_ROOT 构建完整的文件路径
        file_path = os.path.join(settings.MEDIA_ROOT, file_name)
        
        # 保存 PDF 文件
        with open(file_path, 'wb+') as destination:
            for chunk in pdf_file.chunks():
                destination.write(chunk)
        
        # PDF 转图片
        pdf_document = fitz.open(file_path)
        image_path = None
        if len(pdf_document) > 0:
            page = pdf_document.load_page(0)
            pix = page.get_pixmap()
            image_name = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}_page_0.png"
            image_path = os.path.join(settings.MEDIA_ROOT, image_name)
            pix.save(image_path)
        
        # 动态设置 schema
        # schema = [entity.strip() for entity in entities if entity.strip()]  # 去除空字符串
        # if not schema:  # 如果用户没选择，默认使用这些
        #     schema = ['姓名', '出生日期', '电话']
        
        # # 初始化信息提取模型
        # ie = Taskflow('information_extraction', schema=schema)

        # # 执行信息提取
        # result = ie({"doc": image_path})
        data = {}
        mes = ocr.ocr(file_path, cls=True)
        mes_str = json.dumps(mes, ensure_ascii=False).replace("\n", "").replace(" ","")
        completion = client.chat.completions.create(
        model="qwen-plus", # 此处以qwen-plus为例，可按需更换模型名称。模型列表：https://help.aliyun.com/zh/model-studio/getting-started/models
        messages=[
            {'role': 'system', 'content': '你是一个简历提取专家，我会给你文字形式的简历，里面还有很多杂乱的数字，你只关心简历的部分，并把它给我提取成对应的json格式返回给我，例如：{姓名：xxx，年龄：xxx}.注意，请只给我简历的部分，不要别的任何东西'},
            {'role': 'user', 'content': mes_str}],
        )
        res = completion.model_dump_json()
        res = json.loads(res)
        cont = res["choices"][0]["message"]["content"]
        # print(cont)
        data = data = json.loads(cont)
        # for key in schema:
        #     if key in result[0]:
        #         data[key] = result[0][key][0]['text']
        data['pdf_url'] = request.build_absolute_uri(f'/media/{image_name}')
        
        # 返回数据
        return Response({
            'code': 200,
            'data': data,
            'message': '解析成功'
        }, status=status.HTTP_200_OK)

class UploadPDFViewAll(APIView):
    parser_classes = [MultiPartParser, FormParser]
    permission_classes = [permissions.AllowAny]

    def post(self, request):
        # 获取上传的 PDF 文件和实体类型
        pdf_file = request.FILES.get('file')
        entities = request.data.get('entities', '').split(',')  # 接收前端传来的实体类型

        if not pdf_file:
            return Response({
                'code': 400,
                'message': '未上传 PDF 文件或文件格式不正确'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 验证文件类型
        allowed_types = ['application/pdf']
        if pdf_file.content_type not in allowed_types:
            return Response({
                'code': 400,
                'message': '不支持的文件类型，请上传 PDF 格式的文件'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 生成唯一的文件名
        import os
        from datetime import datetime
        from django.conf import settings
        
        # file_name = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}{os.path.splitext(pdf_file.name)[1]}"
        file_name = 'test_pdf.pdf'
        
        # 确保 media 目录存在
        os.makedirs(settings.MEDIA_ROOT, exist_ok=True)
        
        # 使用 MEDIA_ROOT 构建完整的文件路径
        file_path = os.path.join(settings.MEDIA_ROOT, file_name)
        
        # 保存 PDF 文件
        with open(file_path, 'wb+') as destination:
            for chunk in pdf_file.chunks():
                destination.write(chunk)
        
        # PDF 转图片
        pdf_document = fitz.open(file_path)
        image_path = None
        if len(pdf_document) > 0:
            page = pdf_document.load_page(0)
            pix = page.get_pixmap()
            image_name = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}_page_0.png"
            image_path = os.path.join(settings.MEDIA_ROOT, image_name)
            pix.save(image_path)
        
        # 动态设置 schema
        schema = [entity.strip() for entity in entities if entity.strip()]  # 去除空字符串
        if not schema:  # 如果用户没选择，默认使用这些
            schema = ['姓名', '出生日期', '电话']
        
        # 初始化信息提取模型
        ie = Taskflow('information_extraction', schema=schema)

        # 执行信息提取
        result = ie({"doc": image_path})
        data = {}
        for key in schema:
            if key in result[0]:
                data[key] = result[0][key][0]['text']
        data['pdf_url'] = request.build_absolute_uri(f'/media/{image_name}')
        
        # 返回数据
        return Response({
            'code': 200,
            'data': data,
            'message': '解析成功'
        }, status=status.HTTP_200_OK)
def get_paragraphs_text(path):
    """
    获取所有段落的文本
    :param path: word路径
    :return: list类型，如：
        ['Test', 'hello world', ...]
    """
    document = Document(path) 
    # 有的简历是表格式样的，因此，不仅需要提取正文，还要提取表格
    col_keys = [] # 获取列名
    col_values = [] # 获取列值
    index_num = 0
    # 表格提取中，需要添加一个去重机制
    fore_str = ""
    cell_text = ""
    for table in document.tables:
        for row_index,row in enumerate(table.rows):
            for col_index,cell in enumerate(row.cells):
                if fore_str != cell.text:
                    if index_num % 2==0:
                        col_keys.append(cell.text)
                    else:
                        col_values.append(cell.text)
                    fore_str = cell.text
                    index_num +=1
                    cell_text += cell.text + '\n'
    # 提取正文文本
    paragraphs_text = ""
    for paragraph in document.paragraphs:
        # 拼接一个list,包括段落的结构和内容
        paragraphs_text += paragraph.text + "\n"
    return cell_text, paragraphs_text

# Word简历上传处理视图
class UploadWordView(APIView):
    parser_classes = [MultiPartParser, FormParser]
    permission_classes = [permissions.AllowAny]

    def post(self, request):
        print('fangwenlllllll')
        # 获取上传的 Word 文件和实体类型
        word_file = request.FILES.get('file')
        entities = request.data.get('entities', '').split(',')  # 接收前端传来的实体类型

        if not word_file:
            return Response({
                'code': 400,
                'message': '未上传 Word 文件或文件格式不正确'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 验证文件类型
        allowed_types = ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']
        if word_file.content_type not in allowed_types:
            return Response({
                'code': 400,
                'message': '不支持的文件类型，请上传 Word 格式的文件'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 生成唯一的文件名
        import os
        from datetime import datetime
        from django.conf import settings
        from docx import Document
        from docx.shared import Inches  
        
        file_name = f"resume_word{os.path.splitext(word_file.name)[1]}"
        
        # 确保 media 目录存在
        os.makedirs(settings.MEDIA_ROOT, exist_ok=True)
        
        # 使用 MEDIA_ROOT 构建完整的文件路径
        file_path = os.path.join(settings.MEDIA_ROOT, file_name)
        
        # 保存 Word 文件
        with open(file_path, 'wb+') as destination:
            for chunk in word_file.chunks():
                destination.write(chunk)
        
        # 动态设置 schema
        schema = [entity.strip() for entity in entities if entity.strip()]  # 去除空字符串
        if not schema:  # 如果用户没选择，默认使用这些
            schema = ['姓名', '出生日期', '电话']
        cell_text, paragraphs_text = get_paragraphs_text(file_path)
        # # 初始化信息提取模型
        # ie = Taskflow('information_extraction', schema=schema)

        # # 执行信息提取
        # result = ie(cell_text, paragraphs_text)
        
        data = {}
        mes = cell_text + paragraphs_text
        mes = mes.replace("\n", "").replace(" ","")
        # print(mes)
        mes_str = json.dumps(mes, ensure_ascii=False)
        completion = client.chat.completions.create(
        model="qwen-plus", # 此处以qwen-plus为例，可按需更换模型名称。模型列表：https://help.aliyun.com/zh/model-studio/getting-started/models
        messages=[
            {'role': 'system', 'content': '你是一个简历提取专家，我会给你文字形式的简历，里面还有很多杂乱的数字，你只关心简历的部分，并把它给我提取成对应的json格式返回给我，例如：{姓名：xxx，年龄：xxx}.注意，请只给我简历的部分，不要别的任何东西，更不要用什么\\n换行符'},
            {'role': 'user', 'content': mes_str}],
        )
        res = completion.model_dump_json()
        res = json.loads(res)
        cont = res["choices"][0]["message"]["content"]
        # print(cont)
        data = json.loads(cont.replace("''' json", "").replace("\n", "").replace(" ",""))
        # print(data)
        # for key in schema:
        #     if key in result[0]:
        #         data[key] = result[0][key][0]['text']
        # data['word_url'] = request.build_absolute_uri(f'/media/{file_name}')
        
        # 返回数据
        return Response({
            'code': 200,
            'data': data,
            'message': '解析成功'
        }, status=status.HTTP_200_OK)

# 多PDF文件上传处理视图
class UploadMultiplePDFView(APIView):
    parser_classes = [MultiPartParser, FormParser]
    permission_classes = [permissions.AllowAny]

    def post(self, request):
        # 获取上传的多个PDF文件和实体类型
        pdf_files = request.FILES.getlist('files')
        print('Received files:', [f.name for f in pdf_files])
        entities = request.data.get('entities', '')
        print('Received entities:', entities)
        entities = entities.split(',') if entities else []

        if not pdf_files:
            return Response({
                'code': 400,
                'message': '未上传PDF文件或文件格式不正确'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 验证文件类型
        allowed_types = ['application/pdf']
        for pdf_file in pdf_files:
            if pdf_file.content_type not in allowed_types:
                return Response({
                    'code': 400,
                    'message': '不支持的文件类型，请上传PDF格式的文件'
                }, status=status.HTTP_400_BAD_REQUEST)
        
        all_ocr_data = {}
        for index, pdf_file in enumerate(pdf_files):
            print('------------')
            print(index)
            # 生成唯一的文件名
            import os
            from datetime import datetime
            from django.conf import settings
            
            file_name = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}{os.path.splitext(pdf_file.name)[1]}"
            
            # 确保media目录存在
            os.makedirs(settings.MEDIA_ROOT, exist_ok=True)
            
            # 使用MEDIA_ROOT构建完整的文件路径
            file_path = os.path.join(settings.MEDIA_ROOT, file_name)
            
            # 保存PDF文件
            with open(file_path, 'wb+') as destination:
                for chunk in pdf_file.chunks():
                    destination.write(chunk)
            
            # # PDF转图片
            # pdf_document = fitz.open(file_path)
            # image_path = None
            # if len(pdf_document) > 0:
            #     page = pdf_document.load_page(0)
            #     pix = page.get_pixmap()
            #     image_name = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}_page_0.png"
            #     image_path = os.path.join(settings.MEDIA_ROOT, image_name)
            #     pix.save(image_path)
            
            mes = ocr.ocr(file_path, cls=True)
            mes_str = json.dumps(mes, ensure_ascii=False).replace(" ","").replace("\n","")
            all_ocr_data[f'pdf_{index}'] = mes_str
        
        need = request.data.get('inputText', '')
        # 处理need变量的逻辑
        conmm = json.dumps(all_ocr_data)
        # 例如：可以将need的值存储到数据库或进行其他处理
        # 假设这里有一个API调用函数，例如send_data_to_api
        completion = client.chat.completions.create(
        model="qwen-plus", # 此处以qwen-plus为例，可按需更换模型名称。模型列表：https://help.aliyun.com/zh/model-studio/getting-started/models
        messages=[
            {'role': 'system', 'content': '你是一个简历筛选专家，我会给你文字形式的简历，里面还有很多杂乱的数字，你只关心简历的部分。我会在开始给你要求，你根据要求把符合要求的简历编号给我，如{1,2,3}.不要别的任何东西'},
            {'role': 'user', 'content': '要求如下：'+need+conmm}],
        )
        res = completion.model_dump_json()
        res = json.loads(res)
        cont = res["choices"][0]["message"]["content"]
        print(cont)
        # data = json.loads(cont)
        # 返回数据
        return Response({
            'code': 200,
            'data': cont,
            'message': '解析成功'
        }, status=status.HTTP_200_OK)
